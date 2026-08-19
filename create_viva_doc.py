import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(10, 25, 47) # Dark Navy #0A192F
        # Add bottom accent line
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(74, 93, 222) # Indigo #4A5DDE
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(217, 119, 6) # Amber #D97706
    return p

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F1F5F9") # Slate Light
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph() # Spacing

def add_callout(doc, title, text, bg_hex="EFF6FF", border_hex="3B82F6"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    r_t = p.add_run(f"📌 {title}\n")
    r_t.bold = True
    r_t.font.size = Pt(11)
    r_t.font.color.rgb = RGBColor(30, 58, 138)
    
    r_body = p.add_run(text)
    r_body.font.size = Pt(10.5)
    r_body.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph()

def build_word_document(output_filepath):
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Title Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_t = p_title.add_run("CAMPUSLINK DATABASE VIVA MASTERY & EXAMINATION GUIDE")
    run_t.bold = True
    run_t.font.size = Pt(22)
    run_t.font.color.rgb = RGBColor(10, 25, 47)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    run_s = p_sub.add_run("Complete Technical & Oral Defense Prep for MySQL / SQLite | UMaT Tarkwa")
    run_s.font.size = Pt(12)
    run_s.font.color.rgb = RGBColor(100, 116, 139)
    
    # Student Metadata Box
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    meta_data = [
        [("Project:", "CampusLink Peer-to-Peer Marketplace"), ("Database Engines:", "MySQL 8.0+ / SQLite 3")],
        [("Institution:", "Univ. of Mines and Technology (UMaT)"), ("Schema Structure:", "9 Normalized Tables (3NF)")]
    ]
    for r_idx, row in enumerate(meta_data):
        for c_idx, (k, v) in enumerate(row):
            cell = meta_table.cell(r_idx, c_idx)
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            r1 = p.add_run(k + " ")
            r1.bold = True
            r1.font.size = Pt(10)
            r2 = p.add_run(v)
            r2.font.size = Pt(10)
            
    doc.add_paragraph()
    
    add_callout(doc, "AI QUIZZING & STUDY INSTRUCTIONS", 
                "This document contains the complete database blueprint, exact SQL DDL/DML, join explanations, aggregate functions, and 25 viva questions with model answers.\n\n"
                "To quiz yourself using an AI model (ChatGPT, Claude, Gemini, etc.):\n"
                "1. Upload or paste this document into the AI.\n"
                "2. Prompt the AI: 'Act as my Database Professor examining me on CampusLink. Ask me 1 question at a time from this document, wait for my response, grade my answer out of 10, explain any missing details, and then ask the next question.'")

    # SECTION 1: ARCHITECTURE OVERVIEW
    add_heading_styled(doc, "1. Database Architecture & ERD Specifications", level=1)
    
    p = doc.add_paragraph()
    p.add_run("CampusLink uses a ").font.size = Pt(11)
    r = p.add_run("9-Table Relational Database Schema")
    r.bold = True
    p.add_run(" designed in Third Normal Form (3NF) to eliminate data redundancy and preserve referential integrity across all peer-to-peer exchanges.")
    
    headers = ["Table Name", "Primary Key", "Foreign Keys", "Core Purpose & Record Types"]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_background(hdr_cells[i], "0A192F")
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
        
    tables_info = [
        ("users", "user_id", "None", "Stores accounts, PBKDF2 hashed passwords, UMaT emails, verification levels, department, hostel."),
        ("categories", "category_id", "None", "Taxonomy lookup table for equipment categories (Surveying, Computing, Mining PPE, Lab tools, etc.)."),
        ("listings", "listing_id", "owner_id -> users\ncategory_id -> categories", "Peer items posted for rent with daily rate, deposit, condition, availability window, and pickup location."),
        ("rental_requests", "request_id", "listing_id -> listings\nborrower_id -> users", "Booking requests submitted by borrowers detailing rental start/end dates, purpose, and approval status."),
        ("rental_transactions", "transaction_id", "request_id -> rental_requests\nlisting_id -> listings\nborrower_id -> users", "Confirmed rental agreements locking gross amount, 10% platform commission, 90% net payout, and deposit."),
        ("maintenance", "maintenance_id", "listing_id -> listings\nreported_by -> users", "Logs equipment repairs, minor/severe damage reports logged upon return, and maintenance expenses."),
        ("reviews", "review_id", "transaction_id -> transactions\nreviewer_id -> users\nreviewee_id -> users", "Two-way 1–5 star peer ratings and text feedback between lenders and borrowers."),
        ("wishlist", "wishlist_id", "user_id -> users\ncategory_id -> categories", "Stores keyword and category watch notifications for student users."),
        ("saved_listings", "saved_id", "user_id -> users\nlisting_id -> listings", "Junction table storing student bookmarks for quick access to marketplace items.")
    ]
    
    for r_idx, (name, pk, fk, desc) in enumerate(tables_info):
        row_cells = table.add_row().cells
        for i, val in enumerate([name, pk, fk, desc]):
            set_cell_background(row_cells[i], "F8FAFC" if r_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(row_cells[i], top=60, bottom=60, left=80, right=80)
            p = row_cells[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if i == 0:
                r.bold = True
                
    doc.add_paragraph()

    # SECTION 2: DDL & CONSTRAINTS
    add_heading_styled(doc, "2. Data Definition Language (DDL) & Constraint Mastery", level=1)
    
    p = doc.add_paragraph()
    p.add_run("Data Definition Language (DDL) consists of SQL statements used to define, alter, or drop database structures. Below are the exact DDL statements used to create the CampusLink database in MySQL and SQLite.").font.size = Pt(11)

    add_heading_styled(doc, "A. Database Creation & Charset Settings (MySQL)", level=2)
    add_code_block(doc, 
"""CREATE DATABASE IF NOT EXISTS campuslink_umat
  DEFAULT CHARACTER SET utf8mb4
  DEFAULT COLLATE utf8mb4_unicode_ci;

USE campuslink_umat;""")

    add_callout(doc, "VIVA EXPLANATION: utf8mb4 vs utf8",
                "MySQL's standard 'utf8' charset only supports up to 3 bytes per character. 'utf8mb4' supports 4 bytes per character, allowing full Unicode support including international characters, mathematical notation, and emojis. 'utf8mb4_unicode_ci' provides case-insensitive collation for accurate string matching.")

    add_heading_styled(doc, "B. Table Schemas with Complete Constraints", level=2)
    
    add_heading_styled(doc, "1. Users Table Schema:", level=3)
    add_code_block(doc,
"""CREATE TABLE IF NOT EXISTS users (
    user_id INT AUTO_INCREMENT PRIMARY KEY,
    name VARCHAR(255) NOT NULL,
    email VARCHAR(255) UNIQUE NOT NULL,
    password_hash VARCHAR(255) NOT NULL,
    student_id VARCHAR(100) UNIQUE NULL,
    phone VARCHAR(50) NOT NULL,
    verification_level ENUM('Unverified', 'Verified Student', 'Verified Staff', 'Admin') NOT NULL DEFAULT 'Unverified',
    account_status ENUM('Active', 'Suspended', 'Pending Verification') NOT NULL DEFAULT 'Active',
    department VARCHAR(150) NOT NULL,
    hostel VARCHAR(150) NULL,
    last_login DATETIME NULL,
    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT chk_email_umat CHECK (email LIKE '%@%')
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;""")

    add_heading_styled(doc, "2. Listings Table Schema (Foreign Keys & Cascade Rules):", level=3)
    add_code_block(doc,
"""CREATE TABLE IF NOT EXISTS listings (
    listing_id INT AUTO_INCREMENT PRIMARY KEY,
    owner_id INT NOT NULL,
    category_id INT NOT NULL,
    title VARCHAR(255) NOT NULL,
    description TEXT NULL,
    subcategory VARCHAR(150) NOT NULL,
    brand VARCHAR(100) NOT NULL,
    model VARCHAR(100) NOT NULL,
    purchase_year INT NULL,
    rental_rate_per_day DECIMAL(10,2) NOT NULL,
    deposit_amount DECIMAL(10,2) NOT NULL,
    `condition` ENUM('New', 'Good', 'Fair', 'Poor') NOT NULL,
    status ENUM('Available', 'Reserved', 'Rented', 'Maintenance', 'Delisted') NOT NULL DEFAULT 'Available',
    pickup_location VARCHAR(255) NOT NULL,
    thumbnail_path VARCHAR(500) NULL,
    available_from DATE NOT NULL,
    available_until DATE NOT NULL,
    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (owner_id) REFERENCES users (user_id) ON DELETE CASCADE,
    FOREIGN KEY (category_id) REFERENCES categories (category_id) ON DELETE CASCADE,
    CONSTRAINT chk_rate_positive CHECK (rental_rate_per_day >= 0),
    CONSTRAINT chk_deposit_positive CHECK (deposit_amount >= 0),
    CONSTRAINT chk_dates_valid CHECK (available_until >= available_from)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;""")

    add_heading_styled(doc, "3. Rental Transactions Table Schema (Financial Log):", level=3)
    add_code_block(doc,
"""CREATE TABLE IF NOT EXISTS rental_transactions (
    transaction_id INT AUTO_INCREMENT PRIMARY KEY,
    request_id INT UNIQUE NOT NULL,
    listing_id INT NOT NULL,
    borrower_id INT NOT NULL,
    rent_start_date DATE NOT NULL,
    rent_end_date DATE NOT NULL,
    actual_return_date DATE NULL,
    total_days INT NOT NULL,
    gross_amount DECIMAL(10,2) NOT NULL,
    commission_amount DECIMAL(10,2) NOT NULL,
    owner_earnings DECIMAL(10,2) NOT NULL,
    deposit_held DECIMAL(10,2) NOT NULL,
    payment_status ENUM('Pending', 'Paid', 'Refunded') NOT NULL DEFAULT 'Pending',
    rental_status ENUM('Active', 'Returned', 'Overdue', 'Cancelled') NOT NULL DEFAULT 'Active',
    return_notes TEXT NULL,
    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (request_id) REFERENCES rental_requests (request_id) ON DELETE CASCADE,
    FOREIGN KEY (listing_id) REFERENCES listings (listing_id) ON DELETE CASCADE,
    FOREIGN KEY (borrower_id) REFERENCES users (user_id) ON DELETE CASCADE,
    CONSTRAINT chk_total_days CHECK (total_days > 0),
    CONSTRAINT chk_gross_positive CHECK (gross_amount >= 0)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;""")

    add_heading_styled(doc, "C. Database Performance Indexes & B-Tree Optimization", level=2)
    p = doc.add_paragraph()
    p.add_run("Indexes are specialized data structures (typically B-Trees) built on table columns to accelerate search and join operations.").font.size = Pt(11)
    
    add_code_block(doc,
"""-- Performance Optimization Indexes in CampusLink
CREATE INDEX idx_listings_owner ON listings(owner_id);
CREATE INDEX idx_listings_category ON listings(category_id);
CREATE INDEX idx_requests_listing ON rental_requests(listing_id);
CREATE INDEX idx_requests_borrower ON rental_requests(borrower_id);
CREATE INDEX idx_transactions_request ON rental_transactions(request_id);
CREATE INDEX idx_reviews_transaction ON reviews(transaction_id);""")

    add_callout(doc, "EXAM EXPLANATION: B-Tree Index Complexity",
                "Without an index, searching for all listings belonging to user_id = 5 requires scanning every row in the listings table (O(N) Full Table Scan). With idx_listings_owner, the database uses a B-Tree index lookup with O(log N) logarithmic time complexity, making queries virtually instantaneous even with millions of rows.")

    # SECTION 3: SQL JOINS DEEP DIVE
    add_heading_styled(doc, "3. Comprehensive SQL Joins Mastery", level=1)
    
    p = doc.add_paragraph()
    p.add_run("SQL Joins are used to combine data from multiple tables based on logical relationships. CampusLink utilizes INNER JOIN, LEFT JOIN, RIGHT JOIN (and RIGHT JOIN emulation) across its 15 business intelligence reports.").font.size = Pt(11)

    add_heading_styled(doc, "A. INNER JOIN (Matching Records in BOTH Tables)", level=2)
    p = doc.add_paragraph()
    p.add_run("An INNER JOIN returns only those rows where the join condition matches in BOTH the left and right tables.").font.size = Pt(11)
    
    add_code_block(doc,
"""-- Query: Retrieve active listings with Owner Name and Category Name
SELECT 
    l.listing_id, 
    l.title, 
    l.rental_rate_per_day, 
    u.name AS owner_name, 
    c.name AS category_name
FROM listings l
INNER JOIN users u ON l.owner_id = u.user_id
INNER JOIN categories c ON l.category_id = c.category_id
WHERE l.status = 'Available'
ORDER BY l.rental_rate_per_day ASC;""")

    add_heading_styled(doc, "B. LEFT JOIN & NULL Checks (Un-borrowed Equipment)", level=2)
    p = doc.add_paragraph()
    p.add_run("A LEFT JOIN returns ALL rows from the left table, and matching rows from the right table. If there is no match, NULL is returned for the right table's columns.").font.size = Pt(11)
    
    add_code_block(doc,
"""-- Report 06 Query: Find listings that have NEVER been borrowed
SELECT 
    l.listing_id, 
    l.title, 
    c.name AS category_name, 
    u.name AS owner_name, 
    l.rental_rate_per_day
FROM listings l
INNER JOIN categories c ON l.category_id = c.category_id
INNER JOIN users u ON l.owner_id = u.user_id
LEFT JOIN rental_transactions t ON l.listing_id = t.listing_id
WHERE t.transaction_id IS NULL;""")

    add_callout(doc, "HOW THIS QUERY WORKS IN CAMPUSLINK",
                "1. 'listings l' is on the left side.\n"
                "2. 'LEFT JOIN rental_transactions t' matches transactions to listings.\n"
                "3. For items never rented, t.transaction_id returns NULL.\n"
                "4. Filtering 'WHERE t.transaction_id IS NULL' isolates un-borrowed equipment.")

    add_heading_styled(doc, "C. RIGHT JOIN & SQLite Emulation", level=2)
    p = doc.add_paragraph()
    p.add_run("A RIGHT JOIN returns ALL rows from the right table and matching rows from the left table.").font.size = Pt(11)
    
    add_code_block(doc,
"""-- Native MySQL RIGHT JOIN Query (All users and their owned items):
SELECT 
    l.title AS owned_item, 
    u.name AS user_name
FROM listings l
RIGHT JOIN users u ON l.owner_id = u.user_id;

-- SQLite RIGHT JOIN Emulation (Swapping Left/Right table order):
SELECT 
    u.name AS user_name, 
    l.title AS owned_item
FROM users u
LEFT JOIN listings l ON u.user_id = l.owner_id;""")

    add_heading_styled(doc, "D. FULL OUTER JOIN Emulation via UNION", level=2)
    add_code_block(doc,
"""-- FULL OUTER JOIN Emulation in MySQL / SQLite
SELECT u.name, l.title FROM users u LEFT JOIN listings l ON u.user_id = l.owner_id
UNION
SELECT u.name, l.title FROM users u RIGHT JOIN listings l ON u.user_id = l.owner_id;""")

    # SECTION 4: AGGREGATES & REPORTS
    add_heading_styled(doc, "4. Aggregate Functions, Grouping & 15 Intelligence Reports", level=1)
    
    p = doc.add_paragraph()
    p.add_run("Aggregate functions (COUNT, SUM, AVG, MIN, MAX) perform calculations on a set of values and return a single summary value. The GROUP BY clause groups rows that have the same values into summary rows.").font.size = Pt(11)

    add_heading_styled(doc, "Key CampusLink Intelligence Report Queries:", level=2)
    
    add_heading_styled(doc, "Report 01: Platform Financial Volume & Commission Summary", level=3)
    add_code_block(doc,
"""SELECT 
    COUNT(transaction_id) AS total_rentals,
    SUM(gross_amount) AS total_gross_volume,
    SUM(commission_amount) AS platform_10_percent_commission,
    SUM(owner_earnings) AS total_lender_payouts
FROM rental_transactions 
WHERE payment_status = 'Paid';""")

    add_heading_styled(doc, "Report 02: Top 10 Earning Lenders (GROUP BY & HAVING)", level=3)
    add_code_block(doc,
"""SELECT 
    u.name AS lender_name,
    u.department,
    u.hostel,
    COUNT(DISTINCT l.listing_id) AS items_listed,
    SUM(t.owner_earnings) AS total_net_earnings
FROM users u
INNER JOIN listings l ON u.user_id = l.owner_id
INNER JOIN rental_transactions t ON l.listing_id = t.listing_id
GROUP BY u.user_id
HAVING SUM(t.owner_earnings) > 0
ORDER BY total_net_earnings DESC;""")

    add_heading_styled(doc, "Report 10: Average Rental Duration per Category", level=3)
    add_code_block(doc,
"""SELECT 
    c.name AS category_name,
    AVG(t.total_days) AS avg_duration_days,
    MAX(t.total_days) AS max_duration_days
FROM categories c
INNER JOIN listings l ON c.category_id = l.category_id
INNER JOIN rental_transactions t ON l.listing_id = t.listing_id
GROUP BY c.category_id;""")

    # SECTION 5: TRANSACTIONS & SECURITY
    add_heading_styled(doc, "5. Database Transactions (ACID) & Security Architecture", level=1)
    
    add_heading_styled(doc, "A. ACID Transaction Implementation", level=2)
    p = doc.add_paragraph()
    p.add_run("ACID stands for Atomicity, Consistency, Isolation, and Durability. In CampusLink, booking approval requires atomic execution across 4 tables:").font.size = Pt(11)
    
    add_code_block(doc,
"""# Python Transaction Logic in controllers.py
conn = db_engine.get_db_connection()
cursor = conn.cursor()
try:
    # 1. Update Request status to Approved
    cursor.execute("UPDATE rental_requests SET status = 'Approved' WHERE request_id = ?;", (req_id,))
    # 2. Update Listing status to Reserved
    cursor.execute("UPDATE listings SET status = 'Reserved' WHERE listing_id = ?;", (listing_id,))
    # 3. Create Transaction record with 10% commission locking
    cursor.execute("INSERT INTO rental_transactions (...) VALUES (...);")
    # 4. Reject conflicting overlapping pending requests
    cursor.execute("UPDATE rental_requests SET status = 'Rejected' WHERE listing_id = ? AND status = 'Pending' AND request_id != ?;", ...)
    
    conn.commit() # ALL STEPS SAVED ATOMICALLY
except Exception as e:
    conn.rollback() # UNDO ALL CHANGES IF ANY STEP FAILS!""")

    add_heading_styled(doc, "B. SQL Injection Protection", level=2)
    p = doc.add_paragraph()
    p.add_run("CampusLink strictly uses Parameterized Prepared Statements with tuple bindings ('?'). User inputs are never concatenated directly into query strings.").font.size = Pt(11)

    # SECTION 6: 25 VIVA QUESTIONS & MODEL ANSWERS
    add_heading_styled(doc, "6. 25 High-Yield Examiner Viva Questions & Model Answers", level=1)
    
    viva_qna = [
        ("Q1: What database engines does CampusLink support and why?",
         "CampusLink supports a dual-engine architecture: SQLite 3 for lightweight local deployment and testing, and MySQL 8.0+ (with InnoDB engine and utf8mb4 charset) for production multi-user scale. All database operations are wrapped in an engine abstraction layer (db_engine.py)."),
        
        ("Q2: How many tables exist in your schema and what normal form is it in?",
         "There are 9 tables in total (users, categories, listings, rental_requests, rental_transactions, maintenance, reviews, wishlist, saved_listings). The schema is normalized to Third Normal Form (3NF) because all non-key attributes are fully dependent on the primary key, and no transitive dependencies exist."),
        
        ("Q3: What is the purpose of AUTO_INCREMENT PRIMARY KEY?",
         "AUTO_INCREMENT automatically generates a unique sequential integer (1, 2, 3...) for every new record inserted into a table, serving as an immutable surrogate Primary Key."),
        
        ("Q4: Explain the difference between PRIMARY KEY and UNIQUE constraint.",
         "A PRIMARY KEY uniquely identifies each row in a table and cannot contain NULL values (a table can have only one Primary Key). A UNIQUE constraint ensures all values in a column are distinct, but allows NULL values unless explicitly defined as NOT NULL (a table can have multiple UNIQUE constraints, e.g. users.email and users.student_id)."),
        
        ("Q5: What is a Foreign Key and why did you use ON DELETE CASCADE?",
         "A Foreign Key establishes a relational constraint between columns in two tables. We used ON DELETE CASCADE (e.g. listings.owner_id referencing users.user_id) so that if a parent record (user) is deleted, all dependent child records (listings, requests, reviews) are deleted automatically, maintaining referential integrity without creating orphan rows."),
        
        ("Q6: Why did you choose DECIMAL(10,2) for financial rates and amounts instead of FLOAT or DOUBLE?",
         "FLOAT and DOUBLE use binary floating-point representation, which causes IEEE 754 rounding inaccuracies (e.g. 0.1 + 0.2 = 0.30000000000000004). DECIMAL(10,2) stores fixed-point exact numbers with 10 digits total and 2 decimal places, guaranteeing 100% financial accuracy for rental rates, deposits, and platform commissions."),
        
        ("Q7: What is an ENUM data type and where is it used in your project?",
         "An ENUM is a string object whose value is chosen from a fixed list of permitted values defined when the table is created. In CampusLink, ENUMs enforce valid states at the database level for verification_level ('Unverified', 'Verified Student', 'Verified Staff', 'Admin'), listing status ('Available', 'Reserved', 'Rented', 'Maintenance', 'Delisted'), and rental_purpose ('Field Trip', 'Final Year Project', etc.)."),
        
        ("Q8: How does CampusLink calculate platform commissions and net earnings?",
         "When a rental request is approved, total_days is calculated as (rent_end_date - rent_start_date) + 1. gross_amount = rental_rate_per_day * total_days. The system calculates a 10% platform fee: commission_amount = gross_amount * 0.10, and owner_earnings = gross_amount - commission_amount (90% payout)."),
        
        ("Q9: What is an Index and which columns are indexed in CampusLink?",
         "An Index is a B-Tree lookup structure that speeds up query execution. We created indexes on high-frequency search and foreign key columns: idx_listings_owner on listings(owner_id), idx_listings_category on listings(category_id), idx_requests_borrower on rental_requests(borrower_id), and idx_reviews_transaction on reviews(transaction_id)."),
        
        ("Q10: How do indexes improve query performance mathematically?",
         "Without an index, searching for a row requires a Full Table Scan with linear time complexity O(N). An index organizes values into a balanced search tree (B-Tree), reducing lookup time complexity to logarithmic O(log N). For 1,000,000 records, a full scan takes up to 1,000,000 checks, while a B-Tree index takes at most ~20 checks."),
        
        ("Q11: Explain INNER JOIN with a concrete query from your project.",
         "An INNER JOIN returns rows only when the join key exists in both tables. In get_filtered_listings, we use 'listings l INNER JOIN users u ON l.owner_id = u.user_id INNER JOIN categories c ON l.category_id = c.category_id' to retrieve listing details along with the owner's name and category name."),
        
        ("Q12: Explain LEFT JOIN and how it was used to find un-borrowed equipment.",
         "A LEFT JOIN returns all rows from the left table regardless of matches in the right table. In Report 06, we execute 'listings l LEFT JOIN rental_transactions t ON l.listing_id = t.listing_id WHERE t.transaction_id IS NULL'. Items that have never been rented produce NULL transaction IDs, isolating un-borrowed listings."),
        
        ("Q13: Does SQLite support RIGHT JOIN natively? How did you handle it?",
         "Older versions of SQLite (before v3.39) do not support native RIGHT JOIN. In SQLite compatibility mode, RIGHT JOIN is emulated by swapping the table order in a LEFT JOIN query ('FROM users u LEFT JOIN listings l ON u.user_id = l.owner_id')."),
        
        ("Q14: How do you emulate FULL OUTER JOIN in MySQL or SQLite?",
         "Since SQLite and older MySQL versions lack a native FULL OUTER JOIN keyword, it is emulated by combining a LEFT JOIN query and a RIGHT JOIN query using the UNION operator, which removes duplicate rows."),
        
        ("Q15: What is the difference between WHERE and HAVING clauses?",
         "WHERE filters individual rows before grouping occurs. HAVING filters aggregated summary rows after GROUP BY grouping has been applied. For example: 'GROUP BY u.user_id HAVING SUM(t.owner_earnings) > 100.00'."),
        
        ("Q16: How do SQL Wildcards ('LIKE') work in your search feature?",
         "The LIKE operator performs pattern matching using wildcard characters: '%' matches zero or more characters, and '_' matches a single character. CampusLink uses 'LIKE %keyword%' to search item titles, descriptions, brands, and models concurrently."),
        
        ("Q17: How does your database prevent booking overlaps for the same equipment?",
         "CampusLink executes a nested subquery filtering out any listing that has an active transaction ('Active' or 'Overdue') where the requested date range overlaps: 'AND l.listing_id NOT IN (SELECT t.listing_id FROM rental_transactions t WHERE t.rental_status IN ('Active', 'Overdue') AND NOT (t.rent_end_date < req_start OR t.rent_start_date > req_end))'."),
        
        ("Q18: What are CHECK constraints and give two examples from your schema?",
         "CHECK constraints enforce domain integrity by validating that row values satisfy a boolean condition before insertion/update. Examples: 'CONSTRAINT chk_rate_positive CHECK (rental_rate_per_day >= 0)' and 'CONSTRAINT chk_dates_valid CHECK (available_until >= available_from)'."),
        
        ("Q19: What is the purpose of the UNIQUE KEY constraint on saved_listings?",
         "The composite constraint 'UNIQUE KEY uq_saved (user_id, listing_id)' prevents a student from bookmarking the exact same listing multiple times in the database."),
        
        ("Q20: What happens in the database when an item is returned damaged?",
         "When an item is returned with damage: (1) rental_transactions rental_status is updated to 'Returned', (2) listings status is updated to 'Maintenance', and (3) a new maintenance ticket record is inserted into the maintenance table logging the issue description, reported_by user ID, and repair cost."),
        
        ("Q21: How is the user Trust Score calculated from the database?",
         "Trust score is derived dynamically in Python from database aggregate queries: Base Score (50) + (Average Rating * 8) + (Completed Rentals * 0.5) - (Late Return Incidents * 5), capped between [0, 100]."),
        
        ("Q22: What are ACID properties and how are they enforced in your code?",
         "ACID stands for Atomicity, Consistency, Isolation, and Durability. In controllers.py, multi-step operations like approve_request wrap all SQL queries in a transaction block using conn.commit() to save all changes atomically or conn.rollback() to undo all changes if an error occurs."),
        
        ("Q23: How do you prevent SQL Injection attacks in CampusLink?",
         "All SQL queries use Parameterized Prepared Statements with tuple parameter bindings ('?'). User input is passed as separate data parameters, preventing attackers from injecting executable SQL command fragments."),
        
        ("Q24: How are passwords stored securely in the database?",
         "Passwords are saved as PBKDF2-HMAC-SHA256 hashes with 100,000 iterations and a dynamic salt (umat_campuslink_2026). Stored format: 'pbkdf2_sha256$100000$salt$hash_hex'."),
        
        ("Q25: What is the difference between ALTER TABLE in MySQL vs SQLite?",
         "MySQL supports full ALTER TABLE syntax including DROP COLUMN. SQLite historically did not support DROP COLUMN, requiring the 12-Step Table Recreation Pattern: Create temp table -> Copy data -> Drop original table -> Rename temp table.")
    ]

    for q, a in viva_qna:
        p_q = doc.add_paragraph()
        r_q = p_q.add_run(q)
        r_q.bold = True
        r_q.font.size = Pt(11)
        r_q.font.color.rgb = RGBColor(74, 93, 222)
        
        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Inches(0.2)
        r_a = p_a.add_run(a)
        r_a.font.size = Pt(10.5)
        r_a.font.color.rgb = RGBColor(30, 41, 59)
        doc.add_paragraph()

    # Save document
    doc.save(output_filepath)
    print(f"[SUCCESS] Word Document created: {output_filepath}")

if __name__ == "__main__":
    import os
    desktop_path = r"C:\Users\Albert\Desktop\CampusLink_Database_Viva_Mastery_Guide.docx"
    proj_path = r"C:\Users\Albert\Desktop\CAMPUS_LINK\CampusLink_Database_Viva_Mastery_Guide.docx"
    build_word_document(desktop_path)
    build_word_document(proj_path)
